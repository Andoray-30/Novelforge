"""
文档解析器

支持 TXT、MD、DOCX 格式的小说文件解析，包含章节智能分割功能。
"""

import re
from dataclasses import dataclass, field
from enum import Enum
from pathlib import Path
from typing import Optional

from novelforge.core.config import Config


class FileType(str, Enum):
    """文件类型"""
    TXT = "txt"
    MD = "md"
    DOCX = "docx"
    UNKNOWN = "unknown"


@dataclass
class Chapter:
    """章节信息"""
    index: int                    # 章节序号
    title: str                    # 章节标题
    content: str                  # 章节内容
    start_line: int = 0          # 起始行号
    end_line: int = 0            # 结束行号
    word_count: int = 0          # 字数
    
    def __post_init__(self):
        """初始化后计算字数"""
        if self.word_count == 0:
            # 中文字符按1字计算，英文单词按1字计算
            chinese_chars = len(re.findall(r'[\u4e00-\u9fff]', self.content))
            english_words = len(re.findall(r'[a-zA-Z]+', self.content))
            self.word_count = chinese_chars + english_words


@dataclass
class ParsedDocument:
    """解析后的文档"""
    file_type: FileType
    file_path: str
    title: Optional[str] = None          # 文档标题
    author: Optional[str] = None         # 作者
    chapters: list[Chapter] = field(default_factory=list)  # 章节列表
    raw_text: str = ""                   # 原始文本
    total_word_count: int = 0            # 总字数
    
    def __post_init__(self):
        """初始化后计算总字数"""
        if self.total_word_count == 0 and self.chapters:
            self.total_word_count = sum(ch.word_count for ch in self.chapters)


# 章节标题的正则表达式模式
CHAPTER_PATTERNS = [
    # 中文章节：第X章/节/回/卷/部
    r'^第[零一二三四五六七八九十百千万\d]+[章节回卷部篇][\s\S]*$',
    # 带括号的中文章节：【第X章】
    r'^【第[零一二三四五六七八九十百千万\d]+[章节回卷部篇]】[\s\S]*$',
    # 数字章节：1. 标题 或 1、标题
    r'^\d+[\.\、\s][^\n]+$',
    # 英文章节：Chapter X
    r'^[Cc]hapter\s*\d+[\.\s]*[^\n]*$',
    # Part/Section
    r'^[Pp]art\s*\d+[\.\s]*[^\n]*$',
    # 卷/篇
    r'^[卷篇部][\s\d]+[^\n]*$',
    # 序章/终章/尾声
    r'^(序章|引子|楔子|终章|尾声|后记|番外)[\s\S]*$',
]


class DocumentParser:
    """文档解析器"""
    
    def __init__(self, config: Optional[Config] = None):
        """
        初始化解析器
        
        Args:
            config: 配置对象
        """
        self.config = config
        self._compiled_patterns = [
            re.compile(pattern, re.MULTILINE) 
            for pattern in CHAPTER_PATTERNS
        ]
    
    def parse(self, file_path: str) -> ParsedDocument:
        """
        解析文档文件
        
        Args:
            file_path: 文件路径
            
        Returns:
            ParsedDocument: 解析后的文档对象
            
        Raises:
            FileNotFoundError: 文件不存在
            ValueError: 不支持的文件格式
        """
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")
        
        file_type = self.detect_file_type(file_path)
        if file_type == FileType.UNKNOWN:
            raise ValueError(f"不支持的文件格式: {path.suffix}")
        
        # 根据文件类型解析
        if file_type == FileType.DOCX:
            raw_text, chapters = self._parse_docx(file_path)
        else:
            # TXT 和 MD 使用文本解析
            encoding = self._detect_encoding(file_path)
            with open(file_path, 'r', encoding=encoding) as f:
                raw_text = f.read()
            
            if file_type == FileType.MD:
                chapters = self._parse_md(raw_text)
            else:
                chapters = self._parse_txt(raw_text)
        
        # 提取文档元信息
        title, author = self._extract_metadata(raw_text, file_type)
        
        return ParsedDocument(
            file_type=file_type,
            file_path=str(path.absolute()),
            title=title,
            author=author,
            chapters=chapters,
            raw_text=raw_text,
        )
    
    def detect_file_type(self, file_path: str) -> FileType:
        """
        检测文件类型
        
        Args:
            file_path: 文件路径
            
        Returns:
            FileType: 文件类型
        """
        suffix = Path(file_path).suffix.lower()
        type_map = {
            '.txt': FileType.TXT,
            '.md': FileType.MD,
            '.markdown': FileType.MD,
            '.docx': FileType.DOCX,
        }
        return type_map.get(suffix, FileType.UNKNOWN)
    
    def _detect_encoding(self, file_path: str) -> str:
        """
        检测文件编码
        
        Args:
            file_path: 文件路径
            
        Returns:
            str: 编码名称
        """
        # 常见编码列表
        encodings = ['utf-8', 'utf-8-sig', 'gbk', 'gb2312', 'gb18030', 'big5']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    f.read(1024)  # 尝试读取一小段
                return encoding
            except (UnicodeDecodeError, UnicodeError):
                continue
        
        # 默认使用 UTF-8
        return 'utf-8'
    
    def _parse_txt(self, text: str) -> list[Chapter]:
        """
        解析 TXT 格式文本
        
        Args:
            text: 原始文本
            
        Returns:
            list[Chapter]: 章节列表
        """
        return self._detect_chapters(text)
    
    def _parse_md(self, text: str) -> list[Chapter]:
        """
        解析 Markdown 格式文本
        
        Args:
            text: 原始文本
            
        Returns:
            list[Chapter]: 章节列表
        """
        # Markdown 标题作为章节分隔
        lines = text.split('\n')
        chapters = []
        current_title = ""
        current_content_lines = []
        current_start_line = 0
        chapter_index = 0
        
        for i, line in enumerate(lines):
            # 检测 Markdown 标题 (# 开头)
            is_heading = line.startswith('#')
            
            # 检测是否是章节标题
            is_chapter = self._is_chapter_title(line.strip('#').strip())
            
            if is_heading and is_chapter:
                # 保存上一章
                if current_content_lines or current_title:
                    content = '\n'.join(current_content_lines).strip()
                    if content:
                        chapters.append(Chapter(
                            index=chapter_index,
                            title=current_title or f"第{chapter_index + 1}章",
                            content=content,
                            start_line=current_start_line,
                            end_line=i - 1,
                        ))
                        chapter_index += 1
                
                # 开始新章
                current_title = line.strip('#').strip()
                current_content_lines = []
                current_start_line = i
            else:
                current_content_lines.append(line)
        
        # 保存最后一章
        if current_content_lines:
            content = '\n'.join(current_content_lines).strip()
            if content:
                chapters.append(Chapter(
                    index=chapter_index,
                    title=current_title or f"第{chapter_index + 1}章",
                    content=content,
                    start_line=current_start_line,
                    end_line=len(lines) - 1,
                ))
        
        # 如果没有检测到章节，尝试普通章节检测
        if not chapters:
            return self._detect_chapters(text)
        
        return chapters
    
    def _parse_docx(self, file_path: str) -> tuple[str, list[Chapter]]:
        """
        解析 DOCX 格式文件
        
        Args:
            file_path: 文件路径
            
        Returns:
            tuple[str, list[Chapter]]: (原始文本, 章节列表)
        """
        try:
            from docx import Document
        except ImportError:
            raise ImportError(
                "需要安装 python-docx 库来解析 DOCX 文件。"
                "请运行: pip install python-docx"
            )
        
        doc = Document(file_path)
        
        # 提取所有段落文本
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)
        
        raw_text = '\n\n'.join(paragraphs)
        
        # 检测章节
        chapters = self._detect_chapters(raw_text)
        
        return raw_text, chapters
    
    def _detect_chapters(self, text: str) -> list[Chapter]:
        """
        检测章节边界
        
        Args:
            text: 原始文本
            
        Returns:
            list[Chapter]: 章节列表
        """
        lines = text.split('\n')
        chapters = []
        
        # 找出所有章节标题行
        chapter_positions = []
        for i, line in enumerate(lines):
            stripped = line.strip()
            if self._is_chapter_title(stripped):
                chapter_positions.append((i, stripped))
        
        # 如果没有检测到章节标题，将整个文本作为一章
        if not chapter_positions:
            return [Chapter(
                index=0,
                title="全文",
                content=text.strip(),
                start_line=0,
                end_line=len(lines) - 1,
            )]
        
        # 根据章节位置分割内容
        for idx, (pos, title) in enumerate(chapter_positions):
            # 确定章节结束位置
            if idx + 1 < len(chapter_positions):
                end_pos = chapter_positions[idx + 1][0] - 1
            else:
                end_pos = len(lines) - 1
            
            # 提取章节内容（不包含标题行）
            content_lines = lines[pos + 1:end_pos + 1]
            content = '\n'.join(content_lines).strip()
            
            # 清理标题（移除可能的特殊字符）
            clean_title = self._clean_chapter_title(title)
            
            chapters.append(Chapter(
                index=idx,
                title=clean_title,
                content=content,
                start_line=pos,
                end_line=end_pos,
            ))
        
        return chapters
    
    def _is_chapter_title(self, line: str) -> bool:
        """
        判断一行是否是章节标题
        
        Args:
            line: 文本行
            
        Returns:
            bool: 是否是章节标题
        """
        if not line:
            return False
        
        # 行长度限制（章节标题通常不会太长）
        if len(line) > 100:
            return False
        
        # 使用预编译的正则表达式匹配
        for pattern in self._compiled_patterns:
            if pattern.match(line):
                return True
        
        return False
    
    def _clean_chapter_title(self, title: str) -> str:
        """
        清理章节标题
        
        Args:
            title: 原始标题
            
        Returns:
            str: 清理后的标题
        """
        # 移除 Markdown 标题符号
        title = title.lstrip('#').strip()
        # 移除全角括号
        title = re.sub(r'^【(.+)】$', r'\1', title)
        return title
    
    def _extract_metadata(self, text: str, file_type: FileType) -> tuple[Optional[str], Optional[str]]:
        """
        提取文档元信息（标题、作者）
        
        Args:
            text: 原始文本
            file_type: 文件类型
            
        Returns:
            tuple[Optional[str], Optional[str]]: (标题, 作者)
        """
        title = None
        author = None
        
        # 检查前100行
        lines = text.split('\n')[:100]
        
        for line in lines:
            line = line.strip()
            
            # 检测标题
            if not title:
                # 书名号标题
                title_match = re.match(r'^《(.+)》$', line)
                if title_match:
                    title = title_match.group(1)
                    continue
                
                # "书名：" 或 "标题：" 格式
                title_match = re.match(r'^(书名|标题|名称)[：:]\s*(.+)$', line)
                if title_match:
                    title = title_match.group(2).strip()
                    continue
            
            # 检测作者
            if not author:
                author_match = re.match(r'^(作者|著者|Writer|Author)[：:]\s*(.+)$', line, re.IGNORECASE)
                if author_match:
                    author = author_match.group(2).strip()
                    continue
        
        return title, author


def parse_document(file_path: str, config: Optional[Config] = None) -> ParsedDocument:
    """
    便捷函数：解析文档
    
    Args:
        file_path: 文件路径
        config: 配置对象
        
    Returns:
        ParsedDocument: 解析后的文档
    """
    parser = DocumentParser(config)
    return parser.parse(file_path)
